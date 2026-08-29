from __future__ import annotations

from pathlib import Path

from docx import Document

SUPPORTED_EXTS = {
    ".txt": "plain text",
    ".md": "markdown",
    ".markdown": "markdown",
    ".pdf": "pdf",
    ".docx": "word document",
}


def _read_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def _read_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    parts = []
    for page in reader.pages:
        parts.append(page.extract_text() or "")
    return "\n\n".join(parts)


def _read_docx(path: Path) -> str:
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def to_markdown(path: Path) -> str:
    """Convert a supported file into Markdown text. Raises ValueError for
    unsupported extensions."""
    ext = path.suffix.lower()
    if ext not in SUPPORTED_EXTS:
        raise ValueError(f"Unsupported file type: {ext or '<none>'}")
    if ext in {".md", ".markdown"}:
        return _read_txt(path)
    if ext == ".pdf":
        text = _read_pdf(path)
    elif ext == ".docx":
        text = _read_docx(path)
    else:
        text = _read_txt(path)
    if not text.strip():
        return ""
    return f"Source: {path.name}\n\n{text}"
